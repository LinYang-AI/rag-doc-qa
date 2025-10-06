"""
Document ingestion module for RAG system.
Handles PDF, TXT, HTML, and other document formats with modern libraries.
"""

import hashlib
import logging
import mimetypes
from datetime import datetime
from pathlib import Path
from typing import List, Dict, Any, Optional, Union
from dataclasses import dataclass, field
import json

import pypdf
from bs4 import BeautifulSoup
import requests
import chardet
from tqdm import tqdm

from rag_doc_qa.config import UPLOAD_DIR, config

logger = logging.getLogger(__name__)

@dataclass
class Document:
    """Represents a processed document with enhanced metadata."""
    content: str
    metadata: Dict[str, Any] = field(default_factory=dict)
    doc_id: str = ""
    source: str = ""
    timestamp: datetime = field(default_factory=datetime.now)
    
    def __post_init__(self):
        """Generate doc_id if not provided."""
        if not self.doc_id:
            self.doc_id = self._generate_id()
    
    def _generate_id(self) -> str:
        """Generate deterministic document ID."""
        content = f"{self.source}:{self.content[:100]}"
        return hashlib.sha256(content.encode()).hexdigest()[:16]
    
    def to_dict(self) -> Dict[str, Any]:
        """Convert to dictionary for serialization."""
        return {
            "content": self.content,
            "metadata": self.metadata,
            "doc_id": self.doc_id,
            "source": self.source,
            "timestamp": self.timestamp.isoformat()
        }
    
    @classmethod
    def from_dict(cls, data: Dict[str, Any]) -> "Document":
        """Create Document from dictionary."""
        data["timestamp"] = datetime.fromisoformat(data["timestamp"])
        return cls(**data)

class DocumentIngestor:
    """Enhanced document ingestion with support for multiple formats."""
    
    def __init__(self, config_override: Optional[Dict] = None):
        self.config = config_override or {}
        self.processed_docs: List[Document] = []
        self.stats = {
            "total_processed": 0,
            "total_failed": 0,
            "processing_time": 0
        }
    
    def ingest(self, 
               source: Union[str, Path, List[Union[str, Path]]],
               recursive: bool = True,
               show_progress: bool = True) -> List[Document]:
        """
        Universal ingestion method for files, directories, or URLs.
        
        Args:
            source: File path, directory, URL, or list of sources
            recursive: Recursively process directories
            show_progress: Show progress bar
            
        Returns:
            List of processed Document objects
        """
        if isinstance(source, list):
            all_docs = []
            for s in tqdm(source, desc="Processing sources", disable=not show_progress):
                docs = self.ingest(s, recursive, show_progress=False)
                all_docs.extend(docs)
            return all_docs
        
        source = str(source)
        
        # Determine source type and process accordingly
        if source.startswith(("http://", "https://")):
            return self.ingest_url(source)
        
        path = Path(source)
        if path.is_file():
            return self.ingest_file(path)
        elif path.is_dir():
            return self.ingest_directory(path, recursive, show_progress)
        else:
            raise ValueError(f"Invalid source: {source}")
    
    def ingest_file(self, file_path: Path) -> List[Document]:
        """
        Ingest a single file with automatic format detection.
        
        Args:
            file_path: Path to the file
            
        Returns:
            List of Document objects
        """
        file_path = Path(file_path)
        
        if not file_path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")
        
        # Check file size
        file_size_mb = file_path.stat().st_size / (1024 * 1024)
        if file_size_mb > config.max_file_size_mb:
            raise ValueError(f"File too large: {file_size_mb:.2f}MB (max: {config.max_file_size_mb}MB)")
        
        # Detect file type
        suffix = file_path.suffix.lower()
        mime_type, _ = mimetypes.guess_type(str(file_path))
        
        logger.info(f"Ingesting {file_path} (type: {mime_type or suffix})")
        
        try:
            start_time = datetime.now()
            
            if suffix == ".pdf" or mime_type == "application/pdf":
                docs = self._ingest_pdf(file_path)
            elif suffix in [".txt", ".md"] or (mime_type and mime_type.startswith("text/")):
                docs = self._ingest_text(file_path)
            elif suffix in [".html", ".htm"]:
                docs = self._ingest_html(file_path)
            elif suffix == ".docx":
                docs = self._ingest_docx(file_path)
            elif suffix == ".csv":
                docs = self._ingest_csv(file_path)
            else:
                logger.warning(f"Unsupported file type: {suffix}")
                return []
            
            # Update statistics
            elapsed = (datetime.now() - start_time).total_seconds()
            self.stats["total_processed"] += len(docs)
            self.stats["processing_time"] += elapsed
            
            return docs
            
        except Exception as e:
            logger.error(f"Error ingesting {file_path}: {e}")
            self.stats["total_failed"] += 1
            raise
    
    def ingest_directory(self, 
                        dir_path: Path, 
                        recursive: bool = True,
                        show_progress: bool = True) -> List[Document]:
        """
        Ingest all supported files in a directory.
        
        Args:
            dir_path: Path to directory
            recursive: Process subdirectories
            show_progress: Show progress bar
            
        Returns:
            List of all processed documents
        """
        dir_path = Path(dir_path)
        if not dir_path.is_dir():
            raise NotADirectoryError(f"Not a directory: {dir_path}")
        
        # Collect all files
        pattern = "**/*" if recursive else "*"
        files = [f for f in dir_path.glob(pattern) 
                if f.is_file() and f.suffix.lower() in config.allowed_file_types]
        
        all_docs = []
        
        # Process files with progress bar
        for file_path in tqdm(files, desc="Processing files", disable=not show_progress):
            try:
                docs = self.ingest_file(file_path)
                all_docs.extend(docs)
            except Exception as e:
                logger.warning(f"Skipping {file_path}: {e}")
        
        logger.info(f"Ingested {len(all_docs)} documents from {dir_path}")
        return all_docs
    
    def ingest_url(self, url: str, timeout: int = 30) -> List[Document]:
        """
        Ingest content from a URL with improved error handling.
        
        Args:
            url: URL to fetch
            timeout: Request timeout in seconds
            
        Returns:
            List of Document objects
        """
        try:
            headers = {
                "User-Agent": "Mozilla/5.0 (compatible; RAG-Bot/1.0)"
            }
            response = requests.get(url, timeout=timeout, headers=headers)
            response.raise_for_status()
            
            content_type = response.headers.get("content-type", "").lower()
            
            if "html" in content_type:
                soup = BeautifulSoup(response.text, "html.parser")
                content = self._extract_text_from_html(soup)
                title = soup.find("title").string if soup.find("title") else url
            elif "pdf" in content_type:
                # Save temporarily and process as PDF
                temp_path = UPLOAD_DIR / f"temp_{hashlib.md5(url.encode()).hexdigest()}.pdf"
                temp_path.write_bytes(response.content)
                try:
                    return self._ingest_pdf(temp_path)
                finally:
                    temp_path.unlink()
            else:
                content = response.text
                title = url
            
            doc = Document(
                content=content,
                metadata={
                    "source_type": "url",
                    "url": url,
                    "title": title,
                    "content_type": content_type
                },
                source=url
            )
            
            return [doc]
            
        except requests.RequestException as e:
            logger.error(f"Error fetching URL {url}: {e}")
            raise
    
    def _ingest_pdf(self, file_path: Path) -> List[Document]:
        """Extract text from PDF with enhanced metadata."""
        docs = []
        
        with open(file_path, "rb") as file:
            try:
                pdf_reader = pypdf.PdfReader(file)
                
                # Extract document-level metadata
                pdf_metadata = pdf_reader.metadata or {}
                
                for page_num, page in enumerate(pdf_reader.pages, 1):
                    try:
                        text = page.extract_text()
                        
                        if text and text.strip():
                            # Try to extract additional page info
                            page_metadata = {
                                "source_type": "pdf",
                                "file_name": file_path.name,
                                "page": page_num,
                                "total_pages": len(pdf_reader.pages),
                                "title": pdf_metadata.get("/Title", ""),
                                "author": pdf_metadata.get("/Author", ""),
                                "subject": pdf_metadata.get("/Subject", ""),
                            }
                            
                            doc = Document(
                                content=text,
                                metadata=page_metadata,
                                source=str(file_path)
                            )
                            docs.append(doc)
                            
                    except Exception as e:
                        logger.warning(f"Error extracting page {page_num}: {e}")
                        
            except pypdf.errors.PdfReadError as e:
                logger.error(f"PDF read error for {file_path}: {e}")
                raise ValueError(f"Cannot read PDF: {file_path}")
        
        return docs
    
    def _ingest_text(self, file_path: Path) -> List[Document]:
        """Extract text with encoding detection."""
        # Detect encoding
        with open(file_path, "rb") as file:
            raw_data = file.read()
            detected = chardet.detect(raw_data)
            encoding = detected.get("encoding", "utf-8")
            confidence = detected.get("confidence", 0)
        
        logger.debug(f"Detected encoding: {encoding} (confidence: {confidence:.2f})")
        
        # Read with detected encoding
        try:
            text = raw_data.decode(encoding)
        except UnicodeDecodeError:
            # Fallback to UTF-8 with error handling
            text = raw_data.decode("utf-8", errors="ignore")
            logger.warning(f"Used UTF-8 fallback for {file_path}")
        
        doc = Document(
            content=text,
            metadata={
                "source_type": "text",
                "file_name": file_path.name,
                "encoding": encoding,
                "file_size": len(raw_data)
            },
            source=str(file_path)
        )
        
        return [doc]
    
    def _ingest_html(self, file_path: Path) -> List[Document]:
        """Extract text from HTML with metadata preservation."""
        with open(file_path, "r", encoding="utf-8", errors="ignore") as file:
            soup = BeautifulSoup(file.read(), "html.parser")
        
        content = self._extract_text_from_html(soup)
        
        # Extract metadata
        title = soup.find("title")
        title_text = title.string if title else file_path.stem
        
        meta_description = soup.find("meta", attrs={"name": "description"})
        description = meta_description.get("content", "") if meta_description else ""
        
        doc = Document(
            content=content,
            metadata={
                "source_type": "html",
                "file_name": file_path.name,
                "title": title_text,
                "description": description
            },
            source=str(file_path)
        )
        
        return [doc]
    
    def _ingest_docx(self, file_path: Path) -> List[Document]:
        """Extract text from DOCX files."""
        try:
            import docx
            doc_obj = docx.Document(file_path)
            
            # Extract text from paragraphs
            full_text = []
            for paragraph in doc_obj.paragraphs:
                if paragraph.text.strip():
                    full_text.append(paragraph.text)
            
            # Extract text from tables
            for table in doc_obj.tables:
                for row in table.rows:
                    row_text = " | ".join(cell.text.strip() for cell in row.cells)
                    if row_text.strip():
                        full_text.append(row_text)
            
            content = "\n\n".join(full_text)
            
            doc = Document(
                content=content,
                metadata={
                    "source_type": "docx",
                    "file_name": file_path.name,
                    "num_paragraphs": len(doc_obj.paragraphs),
                    "num_tables": len(doc_obj.tables)
                },
                source=str(file_path)
            )
            
            return [doc]
            
        except ImportError:
            logger.error("python-docx not installed. Install with: pip install python-docx")
            return []
    
    def _ingest_csv(self, file_path: Path) -> List[Document]:
        """Extract text from CSV files."""
        try:
            import pandas as pd
            
            df = pd.read_csv(file_path)
            
            # Convert DataFrame to readable text
            content = f"CSV Data from {file_path.name}\n\n"
            content += f"Columns: {', '.join(df.columns)}\n"
            content += f"Rows: {len(df)}\n\n"
            content += df.to_string()
            
            doc = Document(
                content=content,
                metadata={
                    "source_type": "csv",
                    "file_name": file_path.name,
                    "num_rows": len(df),
                    "num_columns": len(df.columns),
                    "columns": list(df.columns)
                },
                source=str(file_path)
            )
            
            return [doc]
            
        except Exception as e:
            logger.error(f"Error reading CSV {file_path}: {e}")
            return []
    
    def _extract_text_from_html(self, soup: BeautifulSoup) -> str:
        """Extract clean text from BeautifulSoup object."""
        # Remove unwanted tags
        for tag in soup(["script", "style", "meta", "link", "noscript"]):
            tag.decompose()
        
        # Extract text with better formatting
        text_parts = []
        
        # Process different tags appropriately
        for tag in soup.find_all(["p", "h1", "h2", "h3", "h4", "h5", "h6", "li", "div"]):
            text = tag.get_text(strip=True)
            if text:
                text_parts.append(text)
        
        # Join with appropriate spacing
        content = "\n\n".join(text_parts)
        
        # Clean up excessive whitespace
        content = " ".join(content.split())
        
        return content
    
    def save_documents(self, docs: List[Document], output_path: Path):
        """Save processed documents to JSON with compression option."""
        output_path = Path(output_path)
        output_path.parent.mkdir(parents=True, exist_ok=True)
        
        data = [doc.to_dict() for doc in docs]
        
        with open(output_path, "w") as f:
            json.dump(data, f, indent=2, default=str)
        
        logger.info(f"Saved {len(docs)} documents to {output_path}")
    
    @staticmethod
    def load_documents(input_path: Path) -> List[Document]:
        """Load documents from JSON."""
        with open(input_path, "r") as f:
            data = json.load(f)
        
        return [Document.from_dict(item) for item in data]
    
    def get_stats(self) -> Dict[str, Any]:
        """Get ingestion statistics."""
        return {
            **self.stats,
            "avg_processing_time": (
                self.stats["processing_time"] / self.stats["total_processed"]
                if self.stats["total_processed"] > 0 else 0
            )
        }

# CLI support
def main():
    """Command-line interface for document ingestion."""
    import argparse
    
    parser = argparse.ArgumentParser(description="Ingest documents for RAG system")
    parser.add_argument("source", help="File, directory, or URL to ingest")
    parser.add_argument("--output", "-o", help="Output JSON file", default="documents.json")
    parser.add_argument("--recursive", "-r", action="store_true", help="Process directories recursively")
    parser.add_argument("--quiet", "-q", action="store_true", help="Suppress progress bars")
    
    args = parser.parse_args()
    
    ingestor = DocumentIngestor()
    docs = ingestor.ingest(args.source, recursive=args.recursive, show_progress=not args.quiet)
    
    if docs:
        ingestor.save_documents(docs, Path(args.output))
        print(f"Successfully ingested {len(docs)} documents")
        print(f"Stats: {ingestor.get_stats()}")
    else:
        print("No documents were ingested")

if __name__ == "__main__":
    main()